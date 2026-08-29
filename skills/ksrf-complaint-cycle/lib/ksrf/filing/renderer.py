"""DOCX/PDF renderer and deterministic semantic/visual checks."""

from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime, timezone
from hashlib import sha256
from pathlib import Path
import re
import shutil
import subprocess
import tempfile
from typing import Any, Iterable

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from PIL import Image
from pypdf import PdfReader

from .composer import StructuredComplaint


PLACEHOLDER_PATTERNS = (
    re.compile(r"\{\{[^{}]+\}\}"),
    re.compile(r"\[(?:УКАЗАТЬ|ВСТАВИТЬ|ЗАПОЛНИТЬ)[^\]]*\]", re.IGNORECASE),
)


@dataclass(frozen=True)
class RenderedArtifact:
    kind: str
    path: str
    mime_type: str
    size: int
    sha256: str
    renderer: str
    renderer_version: str
    status: str
    page_count: int | None = None

    def to_dict(self) -> dict[str, Any]:
        return {
            "kind": self.kind,
            "path": self.path,
            "mime_type": self.mime_type,
            "size": self.size,
            "sha256": self.sha256,
            "renderer": self.renderer,
            "renderer_version": self.renderer_version,
            "status": self.status,
            "page_count": self.page_count,
        }


def file_sha256(path: Path) -> str:
    digest = sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def normalize_text(value: str) -> str:
    return " ".join((value or "").replace("\u00a0", " ").split())


def complaint_plain_text(complaint: StructuredComplaint) -> str:
    chunks = [complaint.title]
    for section in complaint.sections:
        chunks.append(section.heading)
        chunks.extend(sentence.text for sentence in section.sentences)
    return "\n".join(chunks)


def _set_cell_margins(cell: Any, top: int = 80, start: int = 80, bottom: int = 80, end: int = 80) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _add_page_number(paragraph: Any) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, end))


def _configure_document(document: Document, complaint: StructuredComplaint) -> None:
    for section in document.sections:
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(30)
        section.right_margin = Mm(15)
        _add_page_number(section.footer.paragraphs[0])

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Mm(12.5)
    normal.paragraph_format.widow_control = True

    if "KSRF Heading" not in document.styles:
        heading = document.styles.add_style("KSRF Heading", WD_STYLE_TYPE.PARAGRAPH)
    else:
        heading = document.styles["KSRF Heading"]
    heading.font.name = "Times New Roman"
    heading._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    heading.font.size = Pt(14)
    heading.font.bold = True
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    heading.paragraph_format.keep_with_next = True

    properties = document.core_properties
    properties.title = complaint.title
    properties.subject = f"Matter {complaint.matter_id}; draft {complaint.draft_id}"
    properties.author = "KSRF filing-readiness system"
    properties.keywords = "КС РФ, конституционная жалоба, evidence map"
    stable_time = datetime(2000, 1, 1, tzinfo=timezone.utc)
    properties.created = stable_time
    properties.modified = stable_time


def render_docx(complaint: StructuredComplaint, output_path: str | Path) -> RenderedArtifact:
    """Render a real DOCX artifact from a validated complaint model."""

    destination = Path(output_path)
    destination.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    _configure_document(document, complaint)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Mm(0)
    run = title.add_run(complaint.title)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    for section in complaint.sections:
        document.add_paragraph(section.heading, style="KSRF Heading")
        for sentence in section.sentences:
            paragraph = document.add_paragraph(sentence.text)
            paragraph.style = document.styles["Normal"]

    document.save(destination)
    return RenderedArtifact(
        kind="complaint_docx",
        path=str(destination),
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        size=destination.stat().st_size,
        sha256=file_sha256(destination),
        renderer="python-docx",
        renderer_version="1.2",
        status="complete",
    )


def extract_docx_text(path: str | Path) -> str:
    document = Document(str(path))
    chunks: list[str] = []
    chunks.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text)
    for table in document.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells if cell.text)
    return "\n".join(chunks)


def resolve_executable(name: str, explicit: str | Path | None = None) -> str | None:
    if explicit:
        candidate = Path(explicit)
        if candidate.is_file() and candidate.exists():
            return str(candidate)
        return None
    return shutil.which(name)


def convert_docx_to_pdf(
    docx_path: str | Path,
    pdf_path: str | Path,
    *,
    soffice_path: str | Path | None = None,
    timeout_seconds: int = 60,
) -> RenderedArtifact:
    """Convert through LibreOffice without invoking a shell."""

    source = Path(docx_path).resolve()
    destination = Path(pdf_path).resolve()
    executable = resolve_executable("soffice", soffice_path) or resolve_executable(
        "libreoffice", soffice_path
    )
    if not executable:
        raise RuntimeError("Не найден LibreOffice/soffice для преобразования PDF")
    destination.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="ksrf-pdf-") as tmp:
        completed = subprocess.run(
            [
                executable,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                tmp,
                str(source),
            ],
            check=False,
            capture_output=True,
            text=True,
            timeout=timeout_seconds,
        )
        produced = Path(tmp) / f"{source.stem}.pdf"
        if completed.returncode != 0 or not produced.is_file():
            detail = normalize_text(completed.stderr or completed.stdout)
            raise RuntimeError(f"LibreOffice не создал PDF: {detail or 'unknown error'}")
        shutil.copy2(produced, destination)

    reader = PdfReader(str(destination))
    return RenderedArtifact(
        kind="complaint_pdf",
        path=str(destination),
        mime_type="application/pdf",
        size=destination.stat().st_size,
        sha256=file_sha256(destination),
        renderer="LibreOffice",
        renderer_version=_executable_version(executable),
        status="complete",
        page_count=len(reader.pages),
    )


def _executable_version(executable: str) -> str:
    completed = subprocess.run(
        [executable, "--version"],
        check=False,
        capture_output=True,
        text=True,
        timeout=10,
    )
    return normalize_text(completed.stdout or completed.stderr) or "unknown"


def extract_pdf_text(path: str | Path) -> tuple[str, int]:
    reader = PdfReader(str(path))
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    return text, len(reader.pages)


def extract_pdf_body_text(path: str | Path) -> tuple[str, int]:
    """Extract body text while ignoring renderer-added standalone page numbers."""

    reader = PdfReader(str(path))
    pages: list[str] = []
    for page in reader.pages:
        lines = (page.extract_text() or "").splitlines()
        pages.append("\n".join(line for line in lines if not line.strip().isdigit()))
    return "\n".join(pages), len(reader.pages)


def find_unresolved_placeholders(text: str) -> list[str]:
    found: list[str] = []
    for pattern in PLACEHOLDER_PATTERNS:
        found.extend(match.group(0) for match in pattern.finditer(text or ""))
    return sorted(set(found))


def _page_image_findings(path: Path) -> list[dict[str, Any]]:
    with Image.open(path) as source:
        image = source.convert("L")
        width, height = image.size
        histogram = image.histogram()
        dark_count = sum(histogram[:245])
        coverage = dark_count / max(width * height, 1)
        findings: list[dict[str, Any]] = []
        if coverage < 0.0005:
            findings.append(
                {
                    "code": "unexpected_blank_page",
                    "material": True,
                    "message_ru": "Страница предпросмотра практически пуста.",
                    "action_ru": "Проверить разрыв страницы и повторно сформировать PDF.",
                    "ink_coverage": round(coverage, 6),
                }
            )
        edges = (
            image.crop((0, 0, width, min(3, height))),
            image.crop((0, max(height - 3, 0), width, height)),
            image.crop((0, 0, min(3, width), height)),
            image.crop((max(width - 3, 0), 0, width, height)),
        )
        if any(sum(edge.histogram()[:245]) for edge in edges):
            findings.append(
                {
                    "code": "content_touches_page_edge",
                    "material": True,
                    "message_ru": "Содержимое касается края страницы и может быть обрезано.",
                    "action_ru": "Исправить поля или масштаб и заново проверить предпросмотр.",
                }
            )
        if width < 700 or height < 900:
            findings.append(
                {
                    "code": "preview_resolution_unreadable",
                    "material": True,
                    "message_ru": "Разрешение предпросмотра недостаточно для надежной визуальной проверки.",
                    "action_ru": "Отрисовать PDF с разрешением не ниже 120 dpi.",
                    "width": width,
                    "height": height,
                }
            )

        # Long dark rules are a deterministic proxy for a rendered table grid.
        # Text strokes are far shorter and therefore do not cross these thresholds.
        pixels = list(image.getdata())
        row_counts = [0] * height
        column_counts = [0] * width
        for offset, value in enumerate(pixels):
            if value < 96:
                y, x = divmod(offset, width)
                row_counts[y] += 1
                column_counts[x] += 1

        def longest_row_run(y: int) -> int:
            longest = current = 0
            offset = y * width
            for x in range(width):
                if pixels[offset + x] < 96:
                    current += 1
                    longest = max(longest, current)
                else:
                    current = 0
            return longest

        def longest_column_run(x: int) -> int:
            longest = current = 0
            for y in range(height):
                if pixels[y * width + x] < 96:
                    current += 1
                    longest = max(longest, current)
                else:
                    current = 0
            return longest

        row_run_lengths = [longest_row_run(y) for y in range(height)]
        column_run_lengths = [longest_column_run(x) for x in range(width)]

        def grouped(indices: Iterable[int]) -> list[list[int]]:
            groups: list[list[int]] = []
            for index in indices:
                if not groups or index > groups[-1][-1] + 1:
                    groups.append([index])
                else:
                    groups[-1].append(index)
            return groups

        horizontal_rules = grouped(
            index
            for index, run_length in enumerate(row_run_lengths)
            if run_length >= max(20, int(width * 0.25))
        )
        vertical_rules = grouped(
            index
            for index, run_length in enumerate(column_run_lengths)
            if run_length >= max(20, int(height * 0.12))
        )
        horizontal_count = len(horizontal_rules)
        vertical_count = len(vertical_rules)
        horizontal_coordinates = [max(group, key=lambda index: row_run_lengths[index]) for group in horizontal_rules]
        vertical_coordinates = [max(group, key=lambda index: column_run_lengths[index]) for group in vertical_rules]

        def has_intersection(x: int, y: int) -> bool:
            for check_y in range(max(0, y - 2), min(height, y + 3)):
                row_offset = check_y * width
                for check_x in range(max(0, x - 2), min(width, x + 3)):
                    if pixels[row_offset + check_x] < 96:
                        return True
            return False

        intersection_count = sum(
            has_intersection(x, y)
            for y in horizontal_coordinates
            for x in vertical_coordinates
        )
        expected_intersections = horizontal_count * vertical_count
        table_geometry_observed = (
            intersection_count >= 2
            and (
                (horizontal_count >= 2 and vertical_count >= 1)
                or (horizontal_count >= 1 and vertical_count >= 2)
            )
        )
        broken_grid = table_geometry_observed and (
            horizontal_count < 2
            or vertical_count < 2
            or intersection_count < expected_intersections
        )
        if broken_grid:
            findings.append(
                {
                    "code": "broken_table_grid",
                    "material": True,
                    "message_ru": "В предпросмотре обнаружена неполная сетка таблицы.",
                    "action_ru": "Проверить границы, перенос строк и разрыв таблицы между страницами.",
                    "detected_horizontal_rules": horizontal_count,
                    "detected_vertical_rules": vertical_count,
                    "detected_intersections": intersection_count,
                    "expected_intersections": expected_intersections,
                }
            )
        return findings


def _text_layout_findings(
    *,
    page_number: int,
    text_runs: Iterable[dict[str, Any]],
    page_lines: Iterable[str],
    headings: Iterable[str],
    page_size: tuple[float, float] | None = None,
) -> list[dict[str, Any]]:
    """Return conservative findings from PDF text coordinates and font metadata."""

    findings: list[dict[str, Any]] = []
    occupied: dict[tuple[float, float], dict[str, Any]] = {}
    for run in text_runs:
        text = normalize_text(str(run.get("text") or ""))
        if not text:
            continue
        try:
            font_size = float(run.get("font_size") or 0)
        except (TypeError, ValueError):
            font_size = 0
        font_name = normalize_text(str(run.get("font_name") or ""))
        if font_size < 9.0:
            findings.append(
                {
                    "code": "unreadable_font_size",
                    "material": True,
                    "message_ru": "В PDF найден текст размером менее 9 пунктов.",
                    "action_ru": "Увеличить шрифт и повторно сформировать документ.",
                    "page": page_number,
                    "font_size": round(font_size, 3),
                    "text_sample": text[:80],
                }
            )
        if not font_name:
            findings.append(
                {
                    "code": "font_metadata_missing",
                    "material": True,
                    "message_ru": "Для текстового фрагмента PDF не удалось подтвердить шрифт.",
                    "action_ru": "Встроить или заменить шрифт и заново сформировать PDF.",
                    "page": page_number,
                    "text_sample": text[:80],
                }
            )
        x_value = run.get("x")
        y_value = run.get("y")
        try:
            if x_value is None or y_value is None:
                raise ValueError("missing text coordinate")
            position = (round(float(x_value), 1), round(float(y_value), 1))
        except (TypeError, ValueError):
            continue
        if page_size is not None:
            page_width, page_height = page_size
            # A deliberately conservative width approximation is enough to
            # identify content wholly outside the MediaBox without treating
            # ordinary font-metric variation as clipping.
            longest_line = max((len(line) for line in str(run.get("text") or "").splitlines()), default=1)
            estimated_width = max(font_size * 0.25, font_size * 0.25 * longest_line)
            estimated_height = max(font_size, 1.0)
            x, y = position
            wholly_outside = (
                x >= page_width
                or x + estimated_width <= 0
                or y - estimated_height >= page_height
                or y + estimated_height <= 0
            )
            partially_outside = (
                x < 0
                or x + estimated_width > page_width
                or y - estimated_height < 0
                or y > page_height
            )
            if wholly_outside:
                findings.append(
                    {
                        "code": "text_run_off_page",
                        "material": True,
                        "message_ru": "Текстовый фрагмент целиком расположен вне границ страницы PDF.",
                        "action_ru": "Исправить координаты, поля или плавающий объект и повторно сформировать PDF.",
                        "page": page_number,
                        "position": [x, y],
                        "page_size": [round(page_width, 1), round(page_height, 1)],
                        "text_sample": text[:80],
                    }
                )
            elif partially_outside:
                findings.append(
                    {
                        "code": "text_run_clipped_by_page",
                        "material": True,
                        "message_ru": "Часть текстового фрагмента выходит за границы страницы PDF.",
                        "action_ru": "Исправить поля, отступ или ширину блока и повторно проверить PDF.",
                        "page": page_number,
                        "position": [x, y],
                        "page_size": [round(page_width, 1), round(page_height, 1)],
                        "text_sample": text[:80],
                    }
                )
        previous = occupied.get(position)
        if previous is not None and normalize_text(str(previous.get("text") or "")) != text:
            findings.append(
                {
                    "code": "overlapping_text_runs",
                    "material": True,
                    "message_ru": "Два разных текстовых фрагмента отрисованы в одной координате.",
                    "action_ru": "Проверить наложение абзацев, колонтитулов и плавающих объектов.",
                    "page": page_number,
                    "position": [position[0], position[1]],
                    "text_samples": [normalize_text(str(previous.get("text") or ""))[:60], text[:60]],
                }
            )
        else:
            occupied[position] = run

    normalized_headings = {normalize_text(item) for item in headings if normalize_text(item)}
    meaningful_lines = [
        normalize_text(line)
        for line in page_lines
        if normalize_text(line) and not normalize_text(line).isdigit()
    ]
    if meaningful_lines and meaningful_lines[-1] in normalized_headings:
        findings.append(
            {
                "code": "orphan_heading",
                "material": True,
                "message_ru": "Заголовок остался последней содержательной строкой страницы.",
                "action_ru": "Перенести заголовок вместе минимум с одним абзацем следующего раздела.",
                "page": page_number,
                "heading": meaningful_lines[-1],
            }
        )
    return findings


def _pagination_findings(
    *,
    pdf_page_sizes: Iterable[tuple[float, float]],
    preview_paths: Iterable[Path],
    page_lines: Iterable[Iterable[str]],
) -> list[dict[str, Any]]:
    """Validate one-to-one, ordered and geometrically consistent pagination."""

    sizes = list(pdf_page_sizes)
    previews = list(preview_paths)
    lines_by_page = [list(lines) for lines in page_lines]
    findings: list[dict[str, Any]] = []

    if not sizes:
        findings.append(
            {
                "code": "pdf_has_no_pages",
                "material": True,
                "message_ru": "PDF не содержит страниц.",
                "action_ru": "Повторно сформировать PDF из исходного DOCX.",
            }
        )
    rounded_sizes = {(round(width, 1), round(height, 1)) for width, height in sizes}
    if len(rounded_sizes) > 1:
        findings.append(
            {
                "code": "inconsistent_pdf_page_size",
                "material": True,
                "message_ru": "Страницы PDF имеют разный формат.",
                "action_ru": "Привести все разделы документа к единому размеру страницы.",
                "page_sizes": sorted([list(item) for item in rounded_sizes]),
            }
        )

    preview_dimensions: list[tuple[int, int]] = []
    for path in previews:
        with Image.open(path) as image:
            preview_dimensions.append(image.size)
    if len(set(preview_dimensions)) > 1:
        findings.append(
            {
                "code": "inconsistent_preview_dimensions",
                "material": True,
                "message_ru": "Страницы предпросмотра имеют разные размеры.",
                "action_ru": "Повторно отрисовать все страницы с единым dpi и форматом.",
                "preview_dimensions": [list(item) for item in preview_dimensions],
            }
        )

    observed_numbers: list[int] = []
    for path in previews:
        match = re.fullmatch(r"page-(\d+)\.png", path.name)
        if match:
            observed_numbers.append(int(match.group(1)))
    expected_numbers = list(range(1, len(sizes) + 1))
    if len(previews) != len(sizes) or observed_numbers != expected_numbers:
        findings.append(
            {
                "code": "preview_page_sequence_mismatch",
                "material": True,
                "message_ru": "Нумерация или количество страниц предпросмотра не совпадает с PDF.",
                "action_ru": "Очистить каталог предпросмотра и заново отрисовать весь PDF.",
                "expected": expected_numbers,
                "observed": observed_numbers,
            }
        )

    if len(lines_by_page) != len(sizes):
        findings.append(
            {
                "code": "pdf_text_page_count_mismatch",
                "material": True,
                "message_ru": "Текстовая проверка охватила не все страницы PDF.",
                "action_ru": "Повторить извлечение текста и визуальную проверку PDF.",
                "expected": len(sizes),
                "observed": len(lines_by_page),
            }
        )
    for page_number, lines in enumerate(lines_by_page, start=1):
        standalone_numbers = [normalize_text(line) for line in lines if normalize_text(line).isdigit()]
        if str(page_number) not in standalone_numbers:
            findings.append(
                {
                    "code": "page_number_mismatch",
                    "material": True,
                    "message_ru": "Номер страницы в PDF отсутствует или не соответствует позиции страницы.",
                    "action_ru": "Исправить поле номера страницы в колонтитуле и повторно сформировать PDF.",
                    "page": page_number,
                    "observed_numbers": standalone_numbers,
                }
            )
    return findings


def _pdf_layout_findings(
    pdf_path: str | Path,
    preview_paths: Iterable[Path],
    *,
    headings: Iterable[str],
) -> list[dict[str, Any]]:
    reader = PdfReader(str(pdf_path))
    page_sizes: list[tuple[float, float]] = []
    page_lines: list[list[str]] = []
    findings: list[dict[str, Any]] = []
    for page_number, page in enumerate(reader.pages, start=1):
        page_size = (float(page.mediabox.width), float(page.mediabox.height))
        page_sizes.append(page_size)
        runs: list[dict[str, Any]] = []

        def visitor_text(text: str, cm: Any, tm: Any, font_dict: Any, font_size: Any) -> None:
            if not normalize_text(text):
                return
            font_name = ""
            if font_dict is not None:
                font_name = str(font_dict.get("/BaseFont") or "")
            runs.append(
                {
                    "text": text,
                    "x": float(tm[4]) + float(cm[4]),
                    "y": float(tm[5]) + float(cm[5]),
                    "font_size": font_size,
                    "font_name": font_name,
                }
            )

        extracted = page.extract_text(visitor_text=visitor_text) or ""
        lines = [str(line) for line in extracted.splitlines()]
        page_lines.append(lines)
        findings.extend(
            _text_layout_findings(
                page_number=page_number,
                text_runs=runs,
                page_lines=lines,
                headings=headings,
                page_size=page_size,
            )
        )
    findings.extend(
        _pagination_findings(
            pdf_page_sizes=page_sizes,
            preview_paths=list(preview_paths),
            page_lines=page_lines,
        )
    )
    return findings


def render_pdf_previews(
    pdf_path: str | Path,
    preview_dir: str | Path,
    *,
    pdftoppm_path: str | Path | None = None,
    timeout_seconds: int = 60,
) -> list[Path]:
    executable = resolve_executable("pdftoppm", pdftoppm_path)
    if not executable:
        raise RuntimeError("Не найден pdftoppm для визуальной проверки PDF")
    destination = Path(preview_dir)
    destination.mkdir(parents=True, exist_ok=True)
    prefix = destination / "page"
    completed = subprocess.run(
        [executable, "-png", "-r", "120", str(Path(pdf_path).resolve()), str(prefix)],
        check=False,
        capture_output=True,
        text=True,
        timeout=timeout_seconds,
    )
    def page_number(path: Path) -> tuple[int, str]:
        match = re.fullmatch(r"page-(\d+)\.png", path.name)
        return (int(match.group(1)), path.name) if match else (10**9, path.name)

    pages = sorted(destination.glob("page-*.png"), key=page_number)
    if completed.returncode != 0 or not pages:
        detail = normalize_text(completed.stderr or completed.stdout)
        raise RuntimeError(f"Не удалось отрисовать страницы PDF: {detail or 'unknown error'}")
    return pages


def validate_rendered_pair(
    complaint: StructuredComplaint,
    docx_path: str | Path,
    pdf_path: str | Path,
    *,
    preview_dir: str | Path,
    pdftoppm_path: str | Path | None = None,
) -> dict[str, Any]:
    expected = normalize_text(complaint_plain_text(complaint))
    docx_text = normalize_text(extract_docx_text(docx_path))
    pdf_raw, page_count = extract_pdf_body_text(pdf_path)
    pdf_text = normalize_text(pdf_raw)

    missing_in_docx = expected not in docx_text
    missing_in_pdf = expected not in pdf_text
    placeholders = find_unresolved_placeholders("\n".join((docx_text, pdf_text)))
    pages = render_pdf_previews(
        pdf_path, preview_dir, pdftoppm_path=pdftoppm_path
    )
    visual_findings: list[dict[str, Any]] = []
    for index, page in enumerate(pages, start=1):
        for finding in _page_image_findings(page):
            visual_findings.append({"page": index, **finding})
    visual_findings.extend(
        _pdf_layout_findings(
            pdf_path,
            pages,
            headings=[section.heading for section in complaint.sections],
        )
    )

    pagination_codes = {
        "pdf_has_no_pages",
        "inconsistent_pdf_page_size",
        "inconsistent_preview_dimensions",
        "preview_page_sequence_mismatch",
        "pdf_text_page_count_mismatch",
        "page_number_mismatch",
    }

    checks = {
        "qa_profile": "deterministic_visual_v1",
        "docx_semantic_match": not missing_in_docx,
        "pdf_semantic_match": not missing_in_pdf,
        "page_count": page_count,
        "preview_count": len(pages),
        "placeholders": placeholders,
        "visual_findings": visual_findings,
        "pagination_consistent": not any(
            item.get("code") in pagination_codes for item in visual_findings
        ),
    }
    checks["passed"] = (
        checks["docx_semantic_match"]
        and checks["pdf_semantic_match"]
        and not placeholders
        and not any(item.get("material") for item in visual_findings)
        and page_count == len(pages)
    )
    return checks
